import importlib.metadata


ORACLE_DISTRIBUTION = "python-docx"
ORACLE_VERSION = "1.2.0"
TAGGED_SOURCE_ROOT = (
    "https://raw.githubusercontent.com/python-openxml/python-docx/"
    "v1.2.0/docs/user"
)
TAGGED_SOURCE_PAGES = {
    "documents": f"{TAGGED_SOURCE_ROOT}/documents.rst",
    "quickstart": f"{TAGGED_SOURCE_ROOT}/quickstart.rst",
    "text": f"{TAGGED_SOURCE_ROOT}/text.rst",
}
HELD_ROW_SOURCE_BODY = (
    "row = table.rows[1]\n"
    "row.cells[0].text = 'Foo bar to you.'\n"
    "row.cells[1].text = 'And a hearty foo bar to you too sir!'\n"
)
HELD_ROW_RDOCX_BODY = (
    "row = table.rows[1]\n"
    "row.cells[0].text = 'Foo bar to you.'\n"
    "row = document.tables[0].rows[1]\n"
    "row.cells[1].text = 'And a hearty foo bar to you too sir!'\n"
)

EXPECTED_MANIFEST_IDS = frozenset(
    {
        "documents.opening-a-document",
        "documents.really-opening-a-document",
        "quickstart.opening-a-document",
        "quickstart.adding-a-paragraph",
        "quickstart.adding-a-table",
        "quickstart.table-cell-text",
        "quickstart.held-row-two-cell-assignment",
        "quickstart.table-iteration",
        "quickstart.table-style",
        "quickstart.adding-runs",
        "text.horizontal-alignment",
        "text.indentation",
        "text.line-spacing",
        "text.paragraph-spacing",
        "text.pagination-properties",
        "text.font-name-and-size",
        "text.font-tristate-and-underline",
    }
)

DOCUMENTED_S33_EXAMPLES = (
    {
        "id": "documents.opening-a-document",
        "page": TAGGED_SOURCE_PAGES["documents"],
        "heading": "Opening a document",
        "body": "from docx import Document\n\ndocument = Document()\ndocument.save('test.docx')\n",
        "transformation": "namespace-only",
        "setup": "empty",
        "observation": "files",
        "expected": ("test.docx",),
    },
    {
        "id": "documents.really-opening-a-document",
        "page": TAGGED_SOURCE_PAGES["documents"],
        "heading": "REALLY opening a document",
        "body": "document = Document('existing-document-file.docx')\ndocument.save('new-file-name.docx')\n",
        "transformation": "namespace-only",
        "setup": "existing-document",
        "observation": "paragraph-texts",
        "expected": ("existing",),
    },
    {
        "id": "quickstart.opening-a-document",
        "page": TAGGED_SOURCE_PAGES["quickstart"],
        "heading": "Opening a document",
        "body": "from docx import Document\n\ndocument = Document()\n",
        "transformation": "namespace-only",
        "setup": "empty",
        "observation": "paragraph-texts",
        "expected": (),
    },
    {
        "id": "quickstart.adding-a-paragraph",
        "page": TAGGED_SOURCE_PAGES["quickstart"],
        "heading": "Adding a paragraph",
        "body": "paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')\n",
        "transformation": "namespace-only",
        "setup": "document",
        "observation": "paragraph-texts",
        "expected": ("Lorem ipsum dolor sit amet.",),
    },
    {
        "id": "quickstart.adding-a-table",
        "page": TAGGED_SOURCE_PAGES["quickstart"],
        "heading": "Adding a table",
        "body": "table = document.add_table(rows=2, cols=2)\n",
        "transformation": "namespace-only",
        "setup": "document",
        "observation": "table-shape",
        "expected": ((2, 2),),
    },
    {
        "id": "quickstart.table-cell-text",
        "page": TAGGED_SOURCE_PAGES["quickstart"],
        "heading": "Adding a table",
        "body": "cell = table.cell(0, 1)\ncell.text = 'parrot, possibly dead'\n",
        "transformation": "namespace-only",
        "setup": "table",
        "observation": "table-cells",
        "expected": (("", "parrot, possibly dead"), ("", "")),
    },
    {
        "id": "quickstart.held-row-two-cell-assignment",
        "page": TAGGED_SOURCE_PAGES["quickstart"],
        "heading": "Adding a table",
        "body": (
            "row = table.rows[1]\n"
            "row.cells[0].text = 'Foo bar to you.'\n"
            "row.cells[1].text = 'And a hearty foo bar to you too sir!'\n"
        ),
        "transformation": "namespace-and-held-row-refetch",
        "setup": "table",
        "observation": "table-cells",
        "expected": (
            ("", ""),
            ("Foo bar to you.", "And a hearty foo bar to you too sir!"),
        ),
    },
    {
        "id": "quickstart.table-iteration",
        "page": TAGGED_SOURCE_PAGES["quickstart"],
        "heading": "Adding a table",
        "body": "for row in table.rows:\n    for cell in row.cells:\n        print(cell.text)\n",
        "transformation": "namespace-only",
        "setup": "populated-table",
        "observation": "table-cells",
        "expected": (("alpha", "beta"), ("gamma", "delta")),
    },
    {
        "id": "quickstart.table-style",
        "page": TAGGED_SOURCE_PAGES["quickstart"],
        "heading": "Adding a table",
        "body": "table.style = 'LightShading-Accent1'\n",
        "transformation": "namespace-only",
        "setup": "table",
        "observation": "table-style",
        "expected": "LightShading-Accent1",
    },
    {
        "id": "quickstart.adding-runs",
        "page": TAGGED_SOURCE_PAGES["quickstart"],
        "heading": "Applying bold and italic",
        "body": "paragraph = document.add_paragraph('Lorem ipsum ')\nparagraph.add_run('dolor sit amet.')\n",
        "transformation": "namespace-only",
        "setup": "document",
        "observation": "runs",
        "expected": (("Lorem ipsum ", "dolor sit amet."),),
    },
    {
        "id": "text.horizontal-alignment",
        "page": TAGGED_SOURCE_PAGES["text"],
        "heading": "Horizontal alignment (justification)",
        "body": "from docx.enum.text import WD_ALIGN_PARAGRAPH\nparagraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER\n",
        "transformation": "namespace-only",
        "setup": "paragraph-format",
        "observation": "alignment",
        "expected": 1,
    },
    {
        "id": "text.indentation",
        "page": TAGGED_SOURCE_PAGES["text"],
        "heading": "Indentation",
        "body": "from docx.shared import Inches\nfrom docx.shared import Pt\nparagraph_format.left_indent = Inches(0.5)\nparagraph_format.right_indent = Pt(24)\nparagraph_format.first_line_indent = Inches(-0.25)\n",
        "transformation": "namespace-only",
        "setup": "paragraph-format",
        "observation": "indentation",
        "expected": (457200, 304800, -228600),
    },
    {
        "id": "text.line-spacing",
        "page": TAGGED_SOURCE_PAGES["text"],
        "heading": "Line spacing",
        "body": "paragraph_format.line_spacing = Pt(18)\nparagraph_format.line_spacing = 1.75\n",
        "transformation": "namespace-only",
        "setup": "paragraph-format",
        "observation": "line-spacing",
        "expected": 1.75,
    },
    {
        "id": "text.paragraph-spacing",
        "page": TAGGED_SOURCE_PAGES["text"],
        "heading": "Paragraph spacing",
        "body": "paragraph_format.space_before = Pt(18)\nparagraph_format.space_after = Pt(12)\n",
        "transformation": "namespace-only",
        "setup": "paragraph-format",
        "observation": "spacing",
        "expected": (228600, 152400),
    },
    {
        "id": "text.pagination-properties",
        "page": TAGGED_SOURCE_PAGES["text"],
        "heading": "Pagination properties",
        "body": "paragraph_format.keep_with_next = True\nparagraph_format.page_break_before = False\n",
        "transformation": "namespace-only",
        "setup": "paragraph-format",
        "observation": "pagination",
        "expected": (True, False),
    },
    {
        "id": "text.font-name-and-size",
        "page": TAGGED_SOURCE_PAGES["text"],
        "heading": "Apply character formatting",
        "body": "from docx.shared import Pt\nfont.name = 'Calibri'\nfont.size = Pt(12)\n",
        "transformation": "namespace-only",
        "setup": "font",
        "observation": "font-name-and-size",
        "expected": ("Calibri", 152400),
    },
    {
        "id": "text.font-tristate-and-underline",
        "page": TAGGED_SOURCE_PAGES["text"],
        "heading": "Apply character formatting",
        "body": "font.italic = True\nfont.italic = False\nfont.italic = None\nfont.underline = True\nfont.underline = WD_UNDERLINE.DOT_DASH\n",
        "transformation": "namespace-only",
        "setup": "font",
        "observation": "font-tristate-and-underline",
        "expected": (None, 9),
    },
)


def _assert_oracle_version():
    assert importlib.metadata.version(ORACLE_DISTRIBUTION) == ORACLE_VERSION


def _namespace_only(body, namespace):
    if namespace == "docx":
        return body
    assert namespace == "rdocx"
    transformed = body.replace("from docx ", "from rdocx ")
    transformed = transformed.replace("from docx.", "from rdocx.")
    assert transformed.replace("from rdocx ", "from docx ").replace(
        "from rdocx.", "from docx."
    ) == body
    return transformed


def _example_body(example, namespace):
    transformed = _namespace_only(example["body"], namespace)
    if namespace == "docx" or example["transformation"] == "namespace-only":
        return transformed
    assert example["id"] == "quickstart.held-row-two-cell-assignment"
    assert example["transformation"] == "namespace-and-held-row-refetch"
    assert transformed == HELD_ROW_SOURCE_BODY
    return HELD_ROW_RDOCX_BODY


def _fresh_cell(document, row, column):
    return document.tables[0].rows[row].cells[column]


def _example_context(setup, root):
    import rdocx

    namespace = {
        "Document": rdocx.Document,
        "Pt": rdocx.Pt,
        "WD_UNDERLINE": rdocx.WD_UNDERLINE,
    }
    if setup == "empty":
        return namespace

    document = rdocx.Document()
    namespace["document"] = document
    if setup == "existing-document":
        source = rdocx.Document()
        source.add_paragraph("existing")
        source.save(root / "existing-document-file.docx")
        return namespace
    if setup == "document":
        return namespace

    if setup in ("table", "populated-table"):
        table = document.add_table(rows=2, cols=2)
        namespace["table"] = table
    if setup == "table":
        return namespace
    if setup == "populated-table":
        for row, values in enumerate((("alpha", "beta"), ("gamma", "delta"))):
            for column, value in enumerate(values):
                _fresh_cell(document, row, column).text = value
        namespace["table"] = document.tables[0]
        return namespace

    paragraph = document.add_paragraph("")
    if setup == "paragraph-format":
        namespace["paragraph_format"] = paragraph.paragraph_format
        return namespace
    if setup == "font":
        run = paragraph.add_run("")
        namespace["font"] = run.font
        return namespace
    raise AssertionError(f"unknown example setup: {setup}")


def _table_cells(document):
    return tuple(
        tuple(cell.text for cell in row.cells) for row in document.tables[0].rows
    )


def _example_observation(example, namespace, root):
    observation = example["observation"]
    if observation == "files":
        return tuple(sorted(path.name for path in root.glob("*.docx")))
    document = namespace["document"]
    if observation == "paragraph-texts":
        return tuple(paragraph.text for paragraph in document.paragraphs)
    if observation == "table-shape":
        return tuple(
            (len(table.rows), len(table.rows[0].cells)) for table in document.tables
        )
    if observation == "table-cells":
        return _table_cells(document)
    if observation == "table-style":
        return document.tables[0].style
    if observation == "runs":
        return tuple(
            tuple(run.text for run in paragraph.runs)
            for paragraph in document.paragraphs
        )
    paragraph = document.paragraphs[-1]
    paragraph_format = paragraph.paragraph_format
    if observation == "alignment":
        value = paragraph_format.alignment
        return int(value) if value is not None else None
    if observation == "indentation":
        return (
            int(paragraph_format.left_indent),
            int(paragraph_format.right_indent),
            int(paragraph_format.first_line_indent),
        )
    if observation == "line-spacing":
        return paragraph_format.line_spacing
    if observation == "spacing":
        return (int(paragraph_format.space_before), int(paragraph_format.space_after))
    if observation == "pagination":
        return (
            paragraph_format.keep_with_next,
            paragraph_format.page_break_before,
        )
    font = paragraph.runs[-1].font
    if observation == "font-name-and-size":
        return (font.name, int(font.size))
    if observation == "font-tristate-and-underline":
        underline = font.underline
        return (font.italic, int(underline) if underline is not None else None)
    raise AssertionError(f"unknown example observation: {observation}")


def _optional_int(value):
    return int(value) if value is not None else None


def _optional_enum(value):
    if value is None or isinstance(value, bool):
        return value
    return int(value)


def _line_spacing_record(value):
    if value is None:
        return None
    if isinstance(value, float):
        return ("relative", value)
    return ("length", int(value))


def _paragraph_record(paragraph, oracle):
    paragraph_format = paragraph.paragraph_format
    runs = []
    for run in paragraph.runs:
        font = run.font
        color = font.color.rgb if oracle else font.color
        runs.append(
            (
                run.text,
                font.name,
                _optional_int(font.size),
                tuple(color) if color is not None else None,
                font.bold,
                font.italic,
                _optional_enum(font.underline),
                font.strike,
            )
        )
    return (
        paragraph.text,
        _optional_enum(paragraph_format.alignment),
        _optional_int(paragraph_format.space_before),
        _optional_int(paragraph_format.space_after),
        _optional_int(paragraph_format.left_indent),
        _optional_int(paragraph_format.right_indent),
        _optional_int(paragraph_format.first_line_indent),
        _line_spacing_record(paragraph_format.line_spacing),
        paragraph_format.keep_with_next,
        paragraph_format.keep_together,
        paragraph_format.page_break_before,
        paragraph_format.widow_control,
        tuple(runs),
    )


def _table_style(table, oracle):
    if table.style is None:
        return None
    style = table.style.style_id if oracle else table.style
    return None if oracle and style == "TableNormal" else style


def _document_record(document, oracle):
    paragraphs = tuple(
        _paragraph_record(paragraph, oracle) for paragraph in document.paragraphs
    )
    tables = []
    for table in document.tables:
        rows = []
        for row in table.rows:
            rows.append(
                tuple(
                    (
                        cell.text,
                        _optional_int(cell.width),
                        _optional_enum(cell.vertical_alignment),
                        tuple(
                            _paragraph_record(paragraph, oracle)
                            for paragraph in cell.paragraphs
                        ),
                    )
                    for cell in row.cells
                )
            )
        tables.append(
            (
                _table_style(table, oracle),
                _optional_enum(table.alignment),
                tuple(rows),
            )
        )
    return paragraphs, tuple(tables)


def _author_parity_document(
    Document,
    Inches,
    Pt,
    RGBColor,
    WD_ALIGN_PARAGRAPH,
    WD_UNDERLINE,
    WD_TABLE_ALIGNMENT,
    WD_CELL_VERTICAL_ALIGNMENT,
    set_color,
    source_path,
    path,
):
    document = Document(source_path)
    paragraph = document.add_paragraph("Alpha ")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(3)
    paragraph_format.space_after = Pt(6)
    paragraph_format.left_indent = Inches(0.5)
    paragraph_format.right_indent = Inches(0.25)
    paragraph_format.first_line_indent = Inches(-0.25)
    paragraph_format.line_spacing = Pt(18)
    paragraph_format.keep_with_next = True
    paragraph_format.keep_together = False
    paragraph_format.page_break_before = False
    paragraph_format.widow_control = True
    run = paragraph.add_run("beta")
    font = run.font
    font.name = "Aptos"
    font.size = Pt(12)
    set_color(font, RGBColor(0x12, 0x34, 0x56))
    font.bold = True
    font.italic = False
    font.underline = WD_UNDERLINE.DOT_DASH
    font.strike = True

    second = document.add_paragraph("Second")
    second.paragraph_format.line_spacing = 1.75
    table = document.add_table(rows=2, cols=2)
    table.style = "LightShading-Accent1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, values in enumerate((("one", "two"), ("three", "four"))):
        for column, value in enumerate(values):
            cell = document.tables[0].rows[row].cells[column]
            cell.text = value
            cell = document.tables[0].rows[row].cells[column]
            cell.width = Inches(2)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
    cell = document.tables[0].rows[0].cells[0]
    nested = cell.add_paragraph("nested")
    nested.paragraph_format.space_after = Pt(6)
    document.save(path)


def test_documented_s33_examples_run_with_declared_transformations(
    tmp_path, monkeypatch
):
    _assert_oracle_version()
    assert TAGGED_SOURCE_ROOT == (
        "https://raw.githubusercontent.com/python-openxml/python-docx/"
        "v1.2.0/docs/user"
    )
    assert {example["id"] for example in DOCUMENTED_S33_EXAMPLES} == (
        EXPECTED_MANIFEST_IDS
    )
    assert len(DOCUMENTED_S33_EXAMPLES) == len(EXPECTED_MANIFEST_IDS)
    for example in DOCUMENTED_S33_EXAMPLES:
        section = example["id"].split(".", 1)[0]
        assert example["page"] == TAGGED_SOURCE_PAGES[section]
        expected_transformation = (
            "namespace-and-held-row-refetch"
            if example["id"] == "quickstart.held-row-two-cell-assignment"
            else "namespace-only"
        )
        assert example["transformation"] == expected_transformation

    held_row = next(
        example
        for example in DOCUMENTED_S33_EXAMPLES
        if example["id"] == "quickstart.held-row-two-cell-assignment"
    )
    assert held_row["body"] == HELD_ROW_SOURCE_BODY
    assert _example_body(held_row, "docx") == HELD_ROW_SOURCE_BODY
    assert _example_body(held_row, "rdocx") == HELD_ROW_RDOCX_BODY
    line_spacing = next(
        example
        for example in DOCUMENTED_S33_EXAMPLES
        if example["id"] == "text.line-spacing"
    )
    assert line_spacing["body"] == (
        "paragraph_format.line_spacing = Pt(18)\n"
        "paragraph_format.line_spacing = 1.75\n"
    )

    for example in DOCUMENTED_S33_EXAMPLES:
        root = tmp_path / example["id"]
        root.mkdir()
        monkeypatch.chdir(root)
        namespace = _example_context(example["setup"], root)
        exec(_example_body(example, "rdocx"), namespace)
        assert _example_observation(example, namespace, root) == example["expected"]


def test_rdocx_and_python_docx_round_trip_the_same_normalized_content(tmp_path):
    _assert_oracle_version()

    import rdocx
    from docx import Document as OracleDocument
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT as OracleCellAlignment
    from docx.enum.table import WD_TABLE_ALIGNMENT as OracleTableAlignment
    from docx.enum.text import WD_ALIGN_PARAGRAPH as OracleParagraphAlignment
    from docx.enum.text import WD_UNDERLINE as OracleUnderline
    from docx.shared import Inches as OracleInches
    from docx.shared import Pt as OraclePt
    from docx.shared import RGBColor as OracleRGBColor

    assert int(rdocx.Inches(1)) == int(OracleInches(1)) == 914400
    assert int(rdocx.Pt(12)) == int(OraclePt(12)) == 152400
    assert int(rdocx.WD_ALIGN_PARAGRAPH.CENTER) == int(
        OracleParagraphAlignment.CENTER
    )
    assert int(rdocx.WD_UNDERLINE.DOT_DASH) == int(OracleUnderline.DOT_DASH)
    assert int(rdocx.WD_TABLE_ALIGNMENT.CENTER) == int(OracleTableAlignment.CENTER)
    assert int(rdocx.WD_CELL_VERTICAL_ALIGNMENT.BOTTOM) == int(
        OracleCellAlignment.BOTTOM
    )

    rdocx_path = tmp_path / "rdocx-authored.docx"
    oracle_path = tmp_path / "python-docx-authored.docx"
    source_path = tmp_path / "styles-source.docx"
    OracleDocument().save(source_path)
    _author_parity_document(
        rdocx.Document,
        rdocx.Inches,
        rdocx.Pt,
        rdocx.RGBColor,
        rdocx.WD_ALIGN_PARAGRAPH,
        rdocx.WD_UNDERLINE,
        rdocx.WD_TABLE_ALIGNMENT,
        rdocx.WD_CELL_VERTICAL_ALIGNMENT,
        lambda font, color: setattr(font, "color", color),
        source_path,
        rdocx_path,
    )
    _author_parity_document(
        OracleDocument,
        OracleInches,
        OraclePt,
        OracleRGBColor,
        OracleParagraphAlignment,
        OracleUnderline,
        OracleTableAlignment,
        OracleCellAlignment,
        lambda font, color: setattr(font.color, "rgb", color),
        source_path,
        oracle_path,
    )

    records = []
    for path in (rdocx_path, oracle_path):
        rdocx_record = _document_record(rdocx.Document(path), oracle=False)
        oracle_record = _document_record(OracleDocument(path), oracle=True)
        assert rdocx_record == oracle_record
        assert tuple(
            paragraph[7] for paragraph in rdocx_record[0]
        ) == (("length", 228600), ("relative", 1.75))
        assert rdocx_record[1][0][0] == "LightShading-Accent1"
        records.append(rdocx_record)
    assert records[0] == records[1]
